import streamlit as st
import csv
import docx
import re
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io # Used for in-memory file handling

st.set_page_config(layout="centered", page_title="Transfer Routes Printer")

st.title('Transfer-inator')
st.write('Upload your Travel Plans spreadsheet (.csv) to generate a word doc.')

# --- File Uploader ---
uploaded_file = st.file_uploader("Choose a CSV file", type="csv")

if uploaded_file is not None:
    # Read the uploaded file into a string
    csv_bytes = uploaded_file.getvalue()
    csv_string = csv_bytes.decode('utf-8')

    st.success('File uploaded successfully! Now configure document details.')

    # --- Document Configuration ---
    st.header('Document Settings')
    title_text = st.text_input('Document Main Title', value='MLC Plans')
    subtitle_text = st.text_input('Document Subtitle', value='April 2026')

    # --- Process Button ---
    if st.button('Generate Document'):
        try:
            doc = docx.Document()

            # DOCUMENT SIZE AND MARGINS
            section = doc.sections[0]
            section.page_height = Mm(297)
            section.page_width = Mm(210)
            section.left_margin = Mm(25.4)
            section.right_margin = Mm(25.4)
            section.top_margin = Mm(25.4)
            section.bottom_margin = Mm(25.4)
            section.header_distance = Mm(0)
            section.footer_distance = Mm(0)

            # TITLE AND SUBTITLE
            title = doc.add_paragraph(title_text)
            subtitle = doc.add_paragraph().add_run(subtitle_text)
            doc.add_paragraph().add_run('__________________________________________________')

            # NORMAL TEXT STYLE
            normal_style = doc.styles['Normal']
            normal_font = normal_style.font
            normal_font.name = 'Georgia'
            normal_font.size = Pt(14)

            # HEADING 1 STYLE
            heading_style = doc.styles['Heading 1']
            heading_font = heading_style.font
            heading_font.name = 'Georgia'
            heading_font.size = Pt(24)

            # HEADING 2 STYLE
            heading_style = doc.styles['Heading 2']
            heading_font = heading_style.font
            heading_font.name = 'Georgia'
            heading_font.size = Pt(18)

            # HEADING 3 STYLE
            heading_style = doc.styles['Heading 3']
            heading_font = heading_style.font
            heading_font.name = 'Georgia'
            heading_font.size = Pt(14)

            # OPEN TRANSFER CSV AND WRITE INTO DOCUMENT (adapted for in-memory file)
            f = io.StringIO(csv_string)
            reader = csv.DictReader(f)

            for row in reader:
                # --- PERSON NAME ---
                person_name = row.get('person', 'Unknown Person')
                if not person_name.strip():
                    continue  # Skip empty rows
                person = doc.add_heading(level=1).add_run(person_name)
                person.bold = True
                doc.add_paragraph()

                # --- INSTRUCTIONS HEADER ---
                instructions_header = doc.add_paragraph().add_run('INSTRUCTIONS')
                instructions_header.bold = True
                instructions_header.font.size = Pt(18)

                # --- TICKETS ---
                doc.add_paragraph(row.get('tickets/ instructions', ''))

                # --- OPTIONAL KEYS ---
                if row.get('keys'):
                    doc.add_paragraph(row['keys'])

                # --- DATE PARSING ---
                date = row.get('start date', '')
                if not date:
                    current_day = 'N/A'
                    month = 'Unknown'
                else:
                    date_numbers = re.split(r'[/.]', date)
                    try:
                        month_num = date_numbers[0]
                        day = date_numbers[1] if len(date_numbers) > 1 else 'N/A'
                    except (ValueError, IndexError):
                        day = 'N/A'
                        month_num = 'N/A'

                    months_dict = {
                        '1': 'January', '01': 'January',
                        '2': 'February', '02': 'February',
                        '3': 'March', '03': 'March',
                        '4': 'April', '04': 'April',
                        '5': 'May', '05': 'May',
                        '6': 'June', '06': 'June',
                        '7': 'July', '07': 'July',
                        '8': 'August', '08': 'August',
                        '9': 'September', '09': 'September',
                        '10': 'October',
                        '11': 'November',
                        '12': 'December'
                    }

                    month = months_dict.get(month_num, 'Unknown')
                    current_day = day

                # --- FIRST DAY HEADER ---
                first_day = doc.add_paragraph().add_run(f"{current_day} {month}:")
                first_day.bold = True

                # --- INSTRUCTIONS LOOP ---
                instruction_keys = [
                    'one','two','three','four','five','six','seven',
                    'eight','nine','ten','eleven','twelve','thirteen','fourteen','fifteen'
                ]

                for key in instruction_keys:
                    value = row.get(key, '').strip()

                    if value.lower().startswith('end'):
                        break

                    doc.add_paragraph(value)

                    # Detect sleep → next day
                    if 'sleep' in value.lower():
                        try:
                            current_day = str(int(current_day) + 1)
                        except ValueError:
                            pass

                        next_day = doc.add_paragraph().add_run(f"{current_day} {month}:")
                        next_day.bold = True

                doc.add_paragraph('_' * 50)

            # Save the document to a BytesIO object (in-memory file)
            doc_title_filename = subtitle_text.replace(' ', '_') + '.docx'
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="Download Document",
                data=buffer.getvalue(),
                file_name=doc_title_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.balloons()
        except Exception as e:
            st.error(f"An error occurred while generating the document: {str(e)}")